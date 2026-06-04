
# ============================================================
# AulaPronta Pro v2.5 Final Comercial
# Editor, DOCX, impressão/PDF melhor, upload de logo/foto,
# painel admin, backup e preparação mobile/PWA.
# ============================================================

import os, re, shutil
from pathlib import Path
from datetime import datetime
from flask import request, redirect, flash, send_file
from werkzeug.utils import secure_filename
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

FINAL_VERSION = "v2.5-final-comercial"
UPLOAD_DIR = Path("static/uploads")
BACKUP_DIR = Path("backups")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
BACKUP_DIR.mkdir(parents=True, exist_ok=True)

def final_safe_filename(name):
    name = secure_filename(name or "")
    if not name:
        return ""
    base, ext = os.path.splitext(name)
    ext = ext.lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp", ".gif", ".svg"]:
        return ""
    stamp = datetime.now().strftime("%Y%m%d%H%M%S")
    return f"{base[:40]}_{stamp}{ext}"

def final_save_upload(field):
    f = request.files.get(field)
    if not f or not f.filename:
        return ""
    filename = final_safe_filename(f.filename)
    if not filename:
        flash("Formato de imagem inválido. Use PNG, JPG, WEBP, GIF ou SVG.")
        return ""
    dest = UPLOAD_DIR / filename
    f.save(dest)
    return f"/static/uploads/{filename}"

def final_days_left(value):
    try:
        return (datetime.fromisoformat(value).date() - datetime.now().date()).days
    except Exception:
        return None

def final_license_box(u):
    d = final_days_left(u["valid_until"]) if u else None
    if d is None:
        text = "Sem prazo definido"
        cls = "license-good"
    elif d < 0:
        text = "Licença expirada"
        cls = "license-expired"
    elif d <= 7:
        text = f"Atenção: vence em {d} dia(s)"
        cls = "license-soon"
    else:
        text = f"{d} dia(s) restantes"
        cls = "license-good"
    return f'<div class="final-license {cls}"><b>Plano {esc(u["plan"] or "Premium")}</b><span>Válido até {esc(u["valid_until"] or "---")} - {text}</span></div>'

def final_material_info(text):
    raw = text or ""
    def grab(pat):
        m = re.search(pat, raw)
        return m.group(1).strip() if m else ""
    info = {
        "school": grab(r"ESCOLA:\s*(.*)"),
        "teacher": grab(r"PROFESSOR\(A\):\s*(.*)"),
        "grade": grab(r"TURMA:\s*(.*?)\s+DATA:"),
        "subject": grab(r"DISCIPLINA:\s*(.*)"),
        "date": grab(r"DATA:\s*(.*)"),
        "title": ""
    }
    for ln in [x.strip() for x in raw.splitlines() if x.strip()]:
        if ln.isupper() and len(ln) > 6 and ":" not in ln and "QUESTÕES" not in ln:
            info["title"] = ln.title()
            break
    return info

def final_subject_art_by_text(text, fallback="Material"):
    info = final_material_info(text)
    subject = info.get("subject") or fallback
    try:
        return subject_art(subject)
    except Exception:
        return "/static/subject_art/default.svg"

def final_logo_or_avatar(u):
    return u["avatar"] if u and u["avatar"] else "/static/icons/icon-192.png"

def final_docx_download(mid):
    if Document is None:
        flash("Exportação DOCX precisa da dependência python-docx. Rode: pip install python-docx")
        return redirect(f"/professor/material/{mid}")
    u = user()
    m = q("SELECT * FROM materials WHERE id=? AND user_id=?", (mid, u["id"]), True)
    if not m:
        flash("Material não encontrado.")
        return redirect("/professor/materiais")
    info = final_material_info(m["student"])
    title = info.get("title") or m["title"] or "Material AulaPronta"

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AulaPronta Pro")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)

    doc.add_paragraph(f"Tipo: {m['type']}")
    if info.get("subject"):
        doc.add_paragraph(f"Disciplina: {info['subject']}")
    if info.get("grade"):
        doc.add_paragraph(f"Turma: {info['grade']}")
    doc.add_paragraph(f"Gerado em: {m['created_at']}")
    doc.add_paragraph("")

    doc.add_heading("Folha do aluno", level=1)
    for line in (m["student"] or "").splitlines():
        doc.add_paragraph(line)

    if m["teacher"]:
        doc.add_page_break()
        doc.add_heading("Folha do professor / gabarito", level=1)
        for line in (m["teacher"] or "").splitlines():
            doc.add_paragraph(line)

    out = Path("static/uploads") / f"material_{mid}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(out)
    return send_file(out, as_attachment=True, download_name=f"{secure_filename(title) or 'material'}.docx")

def final_editor(mid):
    u = user()
    m = q("SELECT * FROM materials WHERE id=? AND user_id=?", (mid, u["id"]), True)
    if not m:
        flash("Material não encontrado.")
        return redirect("/professor/materiais")
    if request.method == "POST":
        student = request.form.get("student", "")
        teacher = request.form.get("teacher", "")
        title = request.form.get("title", m["title"])
        q("UPDATE materials SET title=?, student=?, teacher=? WHERE id=? AND user_id=?", (title, student, teacher, mid, u["id"]))
        flash("Material editado e salvo com sucesso.")
        return redirect(f"/professor/material/{mid}")
    content = f"""
    <div class="header">
      <div><span>Editor Premium</span><h1>Editar material antes de exportar</h1>
      <p>Revise o texto, ajuste questões, corrija nomes e deixe pronto para baixar em PDF/Word.</p></div>
    </div>
    <div class="panel">
      <form method="post">
        <p><label>Título do material</label><input name="title" value="{esc(m['title'])}"></p>
        <p><label>Folha do aluno</label><textarea name="student" rows="18">{esc(m['student'])}</textarea></p>
        <p><label>Folha do professor / gabarito</label><textarea name="teacher" rows="14">{esc(m['teacher'])}</textarea></p>
        <div class="action-row">
          <button class="btn primary big">Salvar alterações</button>
          <a class="btn" href="/professor/material/{mid}">Cancelar</a>
        </div>
      </form>
    </div>"""
    return render(content, "Editar material")

def final_material_page(mid):
    u = user()
    m = q("SELECT * FROM materials WHERE id=? AND user_id=?", (mid, u["id"]), True)
    if not m:
        flash("Material não encontrado ou não pertence a esta conta.")
        return redirect("/professor/materiais")
    info = final_material_info(m["student"])
    subject = info.get("subject") or m["type"]
    title = info.get("title") or m["title"]
    img = final_subject_art_by_text(m["student"], subject)
    teacher_html = f'<h2>Folha do professor / gabarito</h2><pre>{safe_pre(m["teacher"])}</pre>' if m["teacher"] else ""
    content = f"""
    <div class="header">
      <div class="material-hero"><img src="{img}" alt="{esc(subject)}">
        <div>
          <span>{esc(m['type'])}</span>
          <h1>{esc(title)}</h1>
          <p>Revise, edite e exporte o material final.</p>
          <div class="material-meta">
            <span class="subject-badge">{esc(subject)}</span>
            <span class="subject-badge">{esc(info.get('grade') or 'Turma')}</span>
            <span class="subject-badge">PDF / Word / Editor</span>
          </div>
        </div>
      </div>
      <div class="header-meta">
        <a class="btn" href="/professor/material/{mid}/editar">Editar antes de exportar</a>
        <a class="btn primary" target="_blank" href="/professor/material/{mid}/print">Salvar PDF</a>
        <a class="btn primary" href="/professor/material/{mid}/docx">Baixar Word</a>
      </div>
    </div>
    <div class="panel">
      <h2>Folha do aluno</h2>
      <pre>{safe_pre(m["student"])}</pre>
      {teacher_html}
    </div>"""
    return render(content, title)

def final_print_material(mid):
    u = user()
    m = q("SELECT * FROM materials WHERE id=? AND user_id=?", (mid, u["id"]), True)
    if not m:
        return "<h1>Material não encontrado</h1>", 404
    info = final_material_info(m["student"])
    subject = info.get("subject") or m["type"]
    title = info.get("title") or m["title"]
    img = final_subject_art_by_text(m["student"], subject)
    logo = final_logo_or_avatar(u)
    teacher_page = f'<section class="page"><h2>Folha do professor / gabarito</h2><pre>{safe_pre(m["teacher"])}</pre></section>' if m["teacher"] else ""
    css = """
    body{font-family:Arial,sans-serif;background:#e9eef5;margin:0;padding:24px;color:#111827}
    .page{background:white;max-width:980px;margin:0 auto 22px;padding:34px;border-radius:18px;box-shadow:0 10px 28px rgba(0,0,0,.10)}
    .top{display:grid;grid-template-columns:86px 1fr 145px;gap:18px;align-items:center;border-bottom:2px solid #e5e7eb;padding-bottom:18px;margin-bottom:18px}
    .top img.logo{width:86px;height:86px;object-fit:cover;border-radius:14px;border:1px solid #d9e0eb}
    .top img.art{width:145px;height:92px;object-fit:contain;border-radius:14px;background:#111827;padding:8px}
    .badge{display:inline-block;background:#111827;color:#d4a84f;padding:7px 11px;border-radius:999px;font-size:12px;font-weight:700;margin-right:7px}
    h1{margin:6px 0 8px;font-size:25px;color:#0f172a}h2{color:#0f172a}p{color:#475569;margin:4px 0}
    pre{white-space:pre-wrap;font-family:Consolas,monospace;line-height:1.58;background:#f8fafc;border:1px solid #e5e7eb;padding:18px;border-radius:14px;color:#111827}
    .printbar{max-width:980px;margin:0 auto 14px;display:flex;justify-content:flex-end;gap:8px}
    .printbar button{background:#111827;color:#fff;border:0;padding:12px 16px;border-radius:10px;font-weight:700}
    @media print{body{padding:0;background:white}.printbar{display:none}.page{box-shadow:none;border-radius:0;max-width:none;margin:0;page-break-after:always}}
    """
    return f"""<html><head><meta charset="utf-8"><title>{esc(title)}</title><style>{css}</style></head>
    <body><div class="printbar"><button onclick="window.print()">Imprimir / Salvar PDF</button></div>
    <section class="page">
      <div class="top">
        <img class="logo" src="{logo}">
        <div><span class="badge">AulaPronta Pro</span><span class="badge">{esc(m['type'])}</span><h1>{esc(title)}</h1><p>{esc(u['school'] or 'Escola')} - {esc(u['name'] or 'Professor(a)')}</p><p>{esc(subject)} - {esc(info.get('grade') or 'Turma')}</p></div>
        <img class="art" src="{img}">
      </div>
      <h2>Folha do aluno</h2>
      <pre>{safe_pre(m['student'])}</pre>
    </section>{teacher_page}</body></html>"""

def final_materiais():
    rows = q("SELECT * FROM materials WHERE user_id=? ORDER BY id DESC", (user()["id"],))
    html = '<div class="header"><div><span>Biblioteca Premium</span><h1>Meus materiais</h1><p>Edite, exporte em PDF ou baixe em Word.</p></div></div><div class="panel">'
    if not rows:
        html += '<p class="muted">Nenhum material gerado ainda.</p>'
    for m in rows:
        info = final_material_info(m["student"])
        img = final_subject_art_by_text(m["student"], m["type"])
        desc = f'{info.get("subject") or m["type"]} - {info.get("grade") or "Turma"} - {m["created_at"]}'
        html += f"""<div class="material-list-card">
          <img src="{img}">
          <div><b>{esc(m['title'])}</b><span>{esc(desc)}</span></div>
          <div class="material-actions">
            <a class="btn" href="/professor/material/{m['id']}">Abrir</a>
            <a class="btn" href="/professor/material/{m['id']}/editar">Editar</a>
            <a class="btn primary" href="/professor/material/{m['id']}/docx">Word</a>
          </div>
        </div>"""
    html += "</div>"
    return render(html, "Meus materiais")

def final_perfil():
    u = user()
    if request.method == "POST":
        avatar = final_save_upload("avatar_file") or request.form.get("avatar", "")
        q("UPDATE users SET name=?, school=?, city=?, state=?, avatar=?, teaching_style=?, local_context=?, default_instructions=? WHERE id=?",
          (request.form.get("name",""), request.form.get("school",""), request.form.get("city",""), request.form.get("state","")[:2], avatar,
           request.form.get("teaching_style","Acolhedor e simples"), request.form.get("local_context",""), request.form.get("default_instructions",""), u["id"]))
        flash("Perfil salvo. Os próximos materiais usarão sua identidade.")
        return redirect("/perfil")
    styles = ["Acolhedor e simples","Direto e objetivo","Detalhado e organizado","Lúdico e participativo","EJA contextualizado","Ensino Médio analítico","Interdisciplinar e visual"]
    style_opts = "".join([f'<option {"selected" if (u["teaching_style"] or "") == s else ""}>{s}</option>' for s in styles])
    img = final_logo_or_avatar(u)
    content = f"""
    <div class="header"><div><span>Perfil profissional</span><h1>Minha identidade no material</h1><p>Configure nome, escola, foto/logo e estilo usado nas atividades.</p></div>{final_license_box(u)}</div>
    <div class="panel">
      <form method="post" enctype="multipart/form-data">
        <div class="profile-editor">
          <img src="{img}" alt="imagem atual">
          <div>
            <p><label>Enviar foto/logo da escola</label><input type="file" name="avatar_file" accept="image/*"></p>
            <p><label>Ou usar URL de imagem</label><input name="avatar" value="{esc(u['avatar'])}" placeholder="https://..."></p>
          </div>
        </div>
        <div class="grid">
          <p><label>Nome do professor</label><input name="name" value="{esc(u['name'])}"></p>
          <p><label>Nome da escola</label><input name="school" value="{esc(u['school'])}"></p>
          <p><label>Cidade</label><input name="city" value="{esc(u['city'])}"></p>
          <p><label>UF</label><input name="state" value="{esc(u['state'])}"></p>
        </div>
        <p><label>Estilo dos materiais</label><select name="teaching_style">{style_opts}</select></p>
        <p><label>Contexto da turma/escola</label><textarea name="local_context" rows="4">{esc(u['local_context'])}</textarea></p>
        <p><label>Orientações padrão do professor</label><textarea name="default_instructions" rows="4">{esc(u['default_instructions'])}</textarea></p>
        <button class="btn primary big">Salvar perfil</button>
      </form>
    </div>"""
    return render(content, "Perfil")

def final_admin():
    users = q("SELECT COUNT(*) c FROM users", one=True)["c"]
    serials = q("SELECT COUNT(*) c FROM serials", one=True)["c"]
    questions = q("SELECT COUNT(*) c FROM questions", one=True)["c"]
    topics = q("SELECT COUNT(*) c FROM topics", one=True)["c"]
    mats = q("SELECT COUNT(*) c FROM materials", one=True)["c"]
    expiring = q("SELECT COUNT(*) c FROM users WHERE is_admin=0 AND valid_until <= date('now','+7 day')", one=True)["c"]
    content = f"""
    <div class="header"><div><span>Gestão Premium</span><h1>Painel administrativo comercial</h1><p>Controle usuários, seriais, banco, backup e manutenção.</p></div></div>
    <div class="cards">
      <div class="card"><small>Usuários</small><strong>{users}</strong><p>contas</p></div>
      <div class="card"><small>Seriais</small><strong>{serials}</strong><p>licenças</p></div>
      <div class="card"><small>Materiais</small><strong>{mats}</strong><p>gerados</p></div>
      <div class="card"><small>Vencendo</small><strong>{expiring}</strong><p>em até 7 dias</p></div>
    </div>
    <div class="panel">
      <h2>Ações rápidas</h2>
      <div class="action-row">
        <a class="btn primary" href="/admin/seriais">Gerar seriais</a>
        <a class="btn" href="/admin/usuarios">Usuários</a>
        <a class="btn" href="/admin/banco">Banco pedagógico</a>
        <a class="btn" href="/admin/diagnostico">Diagnóstico</a>
        <a class="btn" href="/admin/backup">Baixar backup</a>
        <a class="btn" href="/admin/checkup">Reconstruir banco</a>
      </div>
    </div>
    <div class="panel"><h2>Resumo do produto</h2><p>Questões: <b>{questions}</b> - Conteúdos: <b>{topics}</b> - Versão: <b>{FINAL_VERSION}</b></p><p class="muted">Antes de vender, troque a senha padrão do administrador e configure uma chave SECRET_KEY forte no RunSite.</p></div>"""
    return render(content, "Admin")

def final_backup():
    db_path = Path("aulapronta.db")
    if not db_path.exists():
        flash("Banco não encontrado.")
        return redirect("/admin")
    out = BACKUP_DIR / f"backup_aulapronta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db"
    shutil.copy2(db_path, out)
    return send_file(out, as_attachment=True, download_name=out.name)

def final_planos_comerciais():
    content = """
    <div class="header"><div><span>AulaPronta Pro</span><h1>Planos comerciais</h1><p>Modelo simples para você vender mensalmente, semestralmente ou anualmente.</p></div></div>
    <div class="cards">
      <div class="card"><small>Mensal</small><strong>R$ 19,90</strong><p>professor individual</p></div>
      <div class="card"><small>Semestral</small><strong>R$ 99,90</strong><p>melhor custo-benefício</p></div>
      <div class="card"><small>Anual</small><strong>R$ 179,90</strong><p>licença completa</p></div>
      <div class="card"><small>Escola</small><strong>sob consulta</strong><p>várias contas</p></div>
    </div>
    <div class="panel"><p>Esta tela é informativa. Para pagamento automático, conecte com Pix/Mercado Pago/Asaas no próximo passo.</p></div>"""
    return render(content, "Planos")

def final_termos():
    return render('<div class="panel"><h1>Termos de uso e privacidade</h1><p>Este sistema auxilia professores na criação de materiais pedagógicos. O professor deve revisar o conteúdo antes de usar em sala de aula.</p><p>Dados de login e materiais são armazenados para funcionamento da plataforma. Configure backup e senha forte antes de vender.</p></div>', "Termos")

app.add_url_rule("/professor/material/<int:mid>/editar", "final_editor", login_required(final_editor), methods=["GET","POST"])
app.add_url_rule("/professor/material/<int:mid>/docx", "final_docx", login_required(final_docx_download))
app.add_url_rule("/admin/backup", "final_backup", admin_required(final_backup))
app.add_url_rule("/planos", "final_planos", final_planos_comerciais)
app.add_url_rule("/termos", "final_termos", final_termos)

app.view_functions["material"] = login_required(final_material_page)
app.view_functions["material_print"] = login_required(final_print_material)
app.view_functions["materiais"] = login_required(final_materiais)
app.view_functions["perfil"] = login_required(final_perfil)
app.view_functions["admin"] = admin_required(final_admin)

try:
    def final_service_worker():
        js = "const CACHE_NAME='aulapronta-pro-v25-final-comercial';const OFFLINE_URL='/offline';self.addEventListener('install',e=>{e.waitUntil(caches.open(CACHE_NAME).then(c=>c.addAll([OFFLINE_URL,'/static/style.css','/static/app.js','/static/login_hero.svg'])))});self.addEventListener('activate',e=>{e.waitUntil(caches.keys().then(keys=>Promise.all(keys.filter(k=>k!==CACHE_NAME).map(k=>caches.delete(k))))) });self.addEventListener('fetch',e=>{if(e.request.method!=='GET')return;e.respondWith(fetch(e.request).catch(()=>caches.match(e.request).then(r=>r||caches.match(OFFLINE_URL))))});"
        return app.response_class(js, mimetype='application/javascript')
    app.view_functions["service_worker"] = final_service_worker
except Exception:
    pass
